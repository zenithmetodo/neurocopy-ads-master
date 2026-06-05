#!/usr/bin/env python3
"""
NeuroCopy Ads — Generador de .docx estéticos para Google Drive
Paleta de marca: Azul (#004195) + Violeta (#9C1AFF) + Turquesa (#00E5CF) + Negro + Blanco

Genera DOS archivos por cada .md:
1. GUIONES_*.docx — Guiones de video (sin "Por qué funciona", con numeración ANUNCIO #)
2. IMAGENES_CREATIVOS_*.docx — Solo creativos estáticos separados
"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === PALETA DE MARCA ===
AZUL = RGBColor(0x00, 0x41, 0x95)              # #004195 - Estructural, confianza
VIOLETA = RGBColor(0x9C, 0x1A, 0xFF)           # #9C1AFF - Acento estratégico
TURQUESA = RGBColor(0x00, 0xE5, 0xCF)          # #00E5CF - CTA, elementos clave
NEGRO = RGBColor(0x00, 0x00, 0x00)             # #000000 - Dramático, elegante
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)            # #FFFFFF - Limpieza, claridad
GRIS_OSCURO = RGBColor(0x2D, 0x2D, 0x2D)      # Texto cuerpo
GRIS_MEDIO = RGBColor(0x5A, 0x5A, 0x5A)       # Texto secundario
TURQUESA_LABEL = RGBColor(0x0D, 0x94, 0x88)   # Turquesa oscuro para labels (contraste)
VIOLETA_OSCURO = RGBColor(0x6D, 0x12, 0xB3)   # Violeta oscuro para labels

# Hex strings for shading (backgrounds)
AZUL_HEX = "004195"
VIOLETA_HEX = "9C1AFF"
TURQUESA_HEX = "00E5CF"
NEGRO_HEX = "000000"

# Fondos claros derivados de la paleta
AZUL_CLARO_HEX = "E8F0FE"         # Fondo [VISUAL] - azul suave
VIOLETA_CLARO_HEX = "F3E8FF"      # Fondo [AUDIO] - violeta suave
TURQUESA_CLARO_HEX = "E0FBF7"     # Fondo [COPY OVERLAY] - turquesa suave
GRIS_CLARO_HEX = "F4F4F5"         # Metadata, fichas
NEGRO_SUAVE_HEX = "1A1A2E"        # Banner ALTO - negro profundo elegante


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    paragraph._p.get_or_add_pPr().append(shading)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_thick_line(doc, color_hex=AZUL_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="24" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_styled_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_level_banner(doc, text, bg_hex=AZUL_HEX, text_color=BLANCO):
    p = doc.add_paragraph()
    set_paragraph_shading(p, bg_hex)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = text_color
    return p


def add_anuncio_banner(doc, anuncio_num, angulo_name=""):
    """Add an ANUNCIO # banner — Turquesa accent for clear client numbering."""
    p = doc.add_paragraph()
    set_paragraph_shading(p, TURQUESA_CLARO_HEX)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    label = f"ANUNCIO {anuncio_num}"
    if angulo_name:
        label += f" — {angulo_name}"
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = AZUL
    return p


def add_structure_badge(doc, structure_name):
    p = doc.add_paragraph()
    set_paragraph_shading(p, VIOLETA_HEX)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"ESTRUCTURA: {structure_name}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLANCO
    return p


def add_metadata_block(doc, lines, bg_hex=GRIS_CLARO_HEX):
    for line in lines:
        if not line.strip():
            continue
        p = doc.add_paragraph()
        set_paragraph_shading(p, bg_hex)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if ":" in line:
            parts = line.split(":", 1)
            run_label = p.add_run(parts[0].strip() + ": ")
            run_label.bold = True
            run_label.font.size = Pt(10)
            run_label.font.color.rgb = AZUL
            run_value = p.add_run(parts[1].strip())
            run_value.font.size = Pt(10)
            run_value.font.color.rgb = GRIS_MEDIO
        else:
            run = p.add_run(line.strip())
            run.font.size = Pt(10)
            run.font.color.rgb = GRIS_MEDIO


def add_hook_header(doc, hook_name):
    p = doc.add_paragraph()
    set_paragraph_shading(p, VIOLETA_CLARO_HEX)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(hook_name)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = VIOLETA


def add_visual_line(doc, text):
    p = doc.add_paragraph()
    set_paragraph_shading(p, AZUL_CLARO_HEX)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run_label = p.add_run("[VISUAL]: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = AZUL
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = GRIS_OSCURO


def add_audio_line(doc, text):
    p = doc.add_paragraph()
    set_paragraph_shading(p, VIOLETA_CLARO_HEX)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run_label = p.add_run("[AUDIO]: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = VIOLETA_OSCURO
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = GRIS_OSCURO


def add_texto_overlay_line(doc, text):
    p = doc.add_paragraph()
    set_paragraph_shading(p, TURQUESA_CLARO_HEX)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run_label = p.add_run("[TEXTO OVERLAY]: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = TURQUESA_LABEL
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = GRIS_OSCURO


def add_section_subheader(doc, text, bg_hex=AZUL_CLARO_HEX, color=AZUL):
    p = doc.add_paragraph()
    set_paragraph_shading(p, bg_hex)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color


def add_notes_header(doc, text="NOTAS DE GRABACIÓN + CONSEJO EDITOR"):
    """Add notes section header — Turquesa accent. Always includes '+ CONSEJO EDITOR'."""
    p = doc.add_paragraph()
    set_paragraph_shading(p, TURQUESA_CLARO_HEX)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    # Force the label to always say "+ CONSEJO EDITOR"
    display_text = text
    if 'CONSEJO' not in display_text.upper():
        display_text = display_text.rstrip() + " + CONSEJO EDITOR"
    run = p.add_run(display_text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = TURQUESA_LABEL


def add_image_creative_header(doc, text):
    p = doc.add_paragraph()
    set_paragraph_shading(p, TURQUESA_CLARO_HEX)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NEGRO


def add_normal_text(doc, text, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRIS_OSCURO
    return p


def add_bold_text(doc, text, color=GRIS_OSCURO):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = color
    return p


def setup_doc_styles(doc):
    """Set up common document styles and margins."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = GRIS_OSCURO
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


def add_field_line(doc, label, value):
    """Add a label: value field line used in metadata and image creatives."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run_label = p.add_run(label + ": ")
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.color.rgb = AZUL
    run_val = p.add_run(value)
    run_val.font.size = Pt(9)
    run_val.font.color.rgb = GRIS_MEDIO


# === FIELDS that belong to image creatives ===
IMAGE_FIELDS = [
    'Nivel de consciencia', 'Angulo', 'Ángulo', 'Deseo que ataca',
    'Tipo visual', 'Call out', 'Texto principal', 'Texto secundario',
    'Descripcion visual', 'Descripción visual', 'CTA en imagen',
    'Dolor que toca', 'Dolor que activa', 'Objecion', 'Objeción'
]

# === FIELDS that belong to ficha técnica ===
FICHA_FIELDS = [
    'Nivel de consciencia', 'Estrategia principal', 'Formato visual',
    'Duracion', 'Duración', 'Angulo', 'Ángulo', 'Deseo que ataca',
    'Dolor que toca', 'Dolor que activa', 'Objecion', 'Objeción',
    'Plataforma', 'Audiencia', 'Call out', 'Tipo visual', 'Tipo:',
    'Texto principal', 'Texto secundario', 'Descripcion visual',
    'CTA en imagen', 'Deseo principal', 'Nivel consciencia',
    'Campo', 'Sub-estructura', 'Fuente original'
]

# === FIELDS to EXCLUDE from client .docx (internal only) ===
CLIENT_EXCLUDE_FIELDS = [
    'Plataforma', 'Fuente original'
]


def is_porque_section(stripped):
    """Check if a line is a POR QUE FUNCIONA header."""
    upper = stripped.upper()
    return 'POR QUE FUNCIONA' in upper or 'POR QUÉ FUNCIONA' in upper


def is_image_section_header(stripped):
    """Check if a line is an image section header."""
    return ('Imágenes' in stripped or 'Imagenes' in stripped or
            'CREATIVOS' in stripped.upper() or 'ESTÁTICOS' in stripped or
            'ESTATICOS' in stripped)


def process_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a styled .docx for CLIENT delivery.

    Key behaviors:
    - SKIPS 'POR QUE FUNCIONA' sections (internal, not for client)
    - SKIPS all IMAGEN CREATIVO blocks (go in separate file)
    - ADDS 'ANUNCIO #' numbering before each guion
    - Renames 'NOTAS DE GRABACION' to 'NOTAS DE GRABACION + CONSEJO EDITOR'
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    setup_doc_styles(doc)

    lines = content.split('\n')
    i = 0
    in_porque = False          # Skipping "por qué funciona" section
    in_image_section = False   # Skipping image creative sections
    anuncio_counter = 0        # ANUNCIO # counter

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            # Empty line can end por-qué or image sections
            if in_porque or in_image_section:
                # Check if the NEXT non-empty line exits the section
                pass
            i += 1
            continue

        # Skip separator lines
        if stripped.startswith('====') or stripped == '---' or stripped == '```':
            if in_porque or in_image_section:
                i += 1
                continue
            i += 1
            continue

        # === POR QUE FUNCIONA: detect and skip entire section ===
        if is_porque_section(stripped):
            in_porque = True
            i += 1
            continue

        # If in "por qué funciona", skip until next major section
        if in_porque:
            # Exit conditions: new level banner, new hook, new guion, new NOTAS, etc.
            if (stripped.startswith('## ') or stripped.startswith('# ') or
                stripped.startswith('### HOOK') or stripped.startswith('### NOTAS') or
                'HOOK' in stripped and stripped.startswith('---') or
                'NIVEL BAJO' in stripped or 'NIVEL MEDIO' in stripped or 'NIVEL ALTO' in stripped or
                'NOTAS DE GRABACION' in stripped.upper() or 'NOTAS DE GRABACIÓN' in stripped.upper() or
                stripped.startswith('GUION') or
                ('IMAGEN CREATIVO' in stripped)):
                in_porque = False
                # Don't skip this line — fall through to process it
            else:
                i += 1
                continue

        # === IMAGE CREATIVES: detect and skip (they go in separate .docx) ===
        if is_image_section_header(stripped) and (stripped.startswith('###') or stripped.startswith('##') or stripped.startswith('#')):
            in_image_section = True
            i += 1
            continue

        if stripped.startswith('IMAGEN CREATIVO') or stripped.startswith('**IMAGEN CREATIVO'):
            in_image_section = True
            i += 1
            continue

        # If in image section, skip until next major non-image section
        if in_image_section:
            if (('NIVEL BAJO' in stripped or 'NIVEL MEDIO' in stripped or 'NIVEL ALTO' in stripped) and
                stripped.startswith('##')):
                in_image_section = False
            elif stripped.startswith('# ') and not is_image_section_header(stripped):
                in_image_section = False
            elif stripped.startswith('GUION') and '#' in stripped:
                in_image_section = False
            else:
                # Check if it's a field inside image creative
                is_img_field = any(stripped.replace('"', '').startswith(f) for f in IMAGE_FIELDS)
                if is_img_field:
                    i += 1
                    continue
                # Other content while in image section: also skip
                i += 1
                continue

        # === MAIN TITLE (# ÁNGULO...) ===
        if stripped.startswith('# ') and not stripped.startswith('## ') and not stripped.startswith('### '):
            title = stripped[2:].strip()
            add_styled_heading(doc, title, level=1, color=AZUL)
            i += 1
            continue

        # === SUBTITLE (## ...) ===
        if stripped.startswith('## ') and not stripped.startswith('### '):
            subtitle = stripped[3:].strip()

            # NIVEL BAJO banner
            if 'NIVEL BAJO' in subtitle and 'ESTRUCTURA:' in subtitle:
                parts = subtitle.split('ESTRUCTURA:')
                level_text = parts[0].strip().rstrip('—').rstrip('-').strip()
                structure = parts[1].strip() if len(parts) > 1 else ""
                anuncio_counter += 1
                add_thick_line(doc, AZUL_HEX)
                add_anuncio_banner(doc, anuncio_counter)
                add_level_banner(doc, level_text, AZUL_HEX, BLANCO)
                if structure:
                    add_structure_badge(doc, structure)
                i += 1
                continue

            # NIVEL MEDIO banner
            elif 'NIVEL MEDIO' in subtitle and 'ESTRUCTURA:' in subtitle:
                parts = subtitle.split('ESTRUCTURA:')
                level_text = parts[0].strip().rstrip('—').rstrip('-').strip()
                structure = parts[1].strip() if len(parts) > 1 else ""
                anuncio_counter += 1
                add_thick_line(doc, VIOLETA_HEX)
                add_anuncio_banner(doc, anuncio_counter)
                add_level_banner(doc, level_text, VIOLETA_HEX, BLANCO)
                if structure:
                    add_structure_badge(doc, structure)
                i += 1
                continue

            # NIVEL ALTO banner
            elif 'NIVEL ALTO' in subtitle and 'ESTRUCTURA:' in subtitle:
                parts = subtitle.split('ESTRUCTURA:')
                level_text = parts[0].strip().rstrip('—').rstrip('-').strip()
                structure = parts[1].strip() if len(parts) > 1 else ""
                anuncio_counter += 1
                add_thick_line(doc, TURQUESA_HEX)
                add_anuncio_banner(doc, anuncio_counter)
                add_level_banner(doc, level_text, NEGRO_SUAVE_HEX, TURQUESA)
                if structure:
                    add_structure_badge(doc, structure)
                i += 1
                continue

            # CREATIVOS/ESTÁTICOS section — skip in guiones file
            elif 'CREATIVOS' in subtitle or 'ESTATICOS' in subtitle or 'ESTÁTICOS' in subtitle:
                in_image_section = True
                i += 1
                continue

            # Retargeting guion headers
            elif 'GUION' in subtitle and ('ÁNGULO:' in subtitle or 'ANGULO:' in subtitle) and 'ESTRUCTURA:' in subtitle:
                parts = subtitle.split('ESTRUCTURA:')
                guion_text = parts[0].strip().rstrip('—').rstrip('-').strip()
                structure = parts[1].strip() if len(parts) > 1 else ""
                anuncio_counter += 1
                add_thick_line(doc, VIOLETA_HEX)
                add_anuncio_banner(doc, anuncio_counter)
                add_level_banner(doc, guion_text, AZUL_HEX, BLANCO)
                if structure:
                    add_structure_badge(doc, structure)
                i += 1
                continue

            else:
                add_styled_heading(doc, subtitle, level=2, color=AZUL)
                i += 1
                continue

        # === SUB-HEADERS (### ...) ===
        if stripped.startswith('### '):
            subheader = stripped[4:].strip()

            if 'ÍNDICE' in subheader or 'INDICE' in subheader:
                add_styled_heading(doc, subheader, level=3, color=AZUL)
                i += 1
                continue

            elif 'Imágenes' in subheader or 'Imagenes' in subheader:
                # Skip — images go to separate file
                in_image_section = True
                i += 1
                continue

            elif 'NOTAS DE GRABACION' in subheader.upper() or 'NOTAS DE GRABACIÓN' in subheader.upper():
                add_notes_header(doc, subheader)
                i += 1
                continue

            elif is_porque_section(subheader):
                in_porque = True
                i += 1
                continue

            elif 'HOOK' in subheader.upper():
                add_hook_header(doc, subheader)
                i += 1
                continue

            elif 'FICHA' in subheader.upper():
                add_section_subheader(doc, subheader, bg_hex=GRIS_CLARO_HEX, color=AZUL)
                i += 1
                continue

            else:
                add_styled_heading(doc, subheader, level=3, color=VIOLETA)
                i += 1
                continue

        # Blockquote metadata — filter out CLIENT_EXCLUDE_FIELDS
        if stripped.startswith('>'):
            meta_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                line_content = lines[i].strip()[1:].strip()
                # Skip excluded fields (Plataforma, Fuente original)
                if not any(line_content.startswith(excl) for excl in CLIENT_EXCLUDE_FIELDS):
                    meta_lines.append(line_content)
                i += 1
            if meta_lines:
                add_metadata_block(doc, meta_lines)
            continue

        # Table rows — filter out CLIENT_EXCLUDE_FIELDS (Plataforma, Fuente original)
        if stripped.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = lines[i].strip()
                if '---' not in row:
                    # Skip rows containing excluded fields
                    row_text = row.replace('*', '').replace('|', ' ')
                    if any(excl in row_text for excl in CLIENT_EXCLUDE_FIELDS):
                        i += 1
                        continue
                    cells = [c.strip() for c in row.split('|')[1:-1]]
                    table_rows.append(cells)
                i += 1

            if len(table_rows) > 1:
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx in range(min(len(row_data), num_cols)):
                        cell = table.rows[row_idx].cells[col_idx]
                        if row_idx == 0:
                            set_cell_shading(cell, AZUL_HEX)
                            p = cell.paragraphs[0]
                            run = p.add_run(row_data[col_idx])
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = BLANCO
                        else:
                            bg = GRIS_CLARO_HEX if row_idx % 2 == 1 else "FFFFFF"
                            set_cell_shading(cell, bg)
                            p = cell.paragraphs[0]
                            run = p.add_run(row_data[col_idx])
                            run.font.size = Pt(9)
                            run.font.color.rgb = GRIS_OSCURO
                            if col_idx == 2 and row_idx > 0:
                                run.bold = True
                                run.font.color.rgb = VIOLETA
            continue

        # CUERPO header
        if 'CUERPO' in stripped and ('---' in stripped or stripped.startswith('**CUERPO')):
            body_name = stripped.replace('---', '').replace('**', '').strip()
            add_hook_header(doc, body_name)
            i += 1
            continue

        # Hook headers
        if ('HOOK' in stripped and ('---' in stripped or stripped.startswith('**HOOK') or stripped.startswith('HOOK'))) or \
           re.match(r'^#{0,3}\s*HOOK\s+[ABC]', stripped):
            hook_name = stripped.replace('---', '').replace('**', '').replace('#', '').strip()
            add_hook_header(doc, hook_name)
            i += 1
            continue

        # [VISUAL] lines
        if '[VISUAL]' in stripped or stripped.startswith('VISUAL:') or stripped.startswith('**VISUAL'):
            text = re.sub(r'^\*{0,2}\[?VISUAL\]?\*{0,2}:?\*{0,2}\s*', '', stripped).strip()
            add_visual_line(doc, text)
            i += 1
            continue

        # [AUDIO] lines
        if '[AUDIO]' in stripped or stripped.startswith('AUDIO:') or stripped.startswith('**AUDIO'):
            text = re.sub(r'^\*{0,2}\[?AUDIO\]?\*{0,2}\s*\(?[^)]*\)?\s*:?\*{0,2}\s*', '', stripped).strip()
            if not text:
                text = re.sub(r'^\*{0,2}\[?AUDIO\]?\*{0,2}:?\*{0,2}\s*', '', stripped).strip()
            add_audio_line(doc, text)
            i += 1
            continue

        # [TEXTO OVERLAY] / [COPY OVERLAY] lines — always render as TEXTO OVERLAY
        if 'COPY OVERLAY' in stripped or 'TEXTO OVERLAY' in stripped:
            text = re.sub(r'^\*{0,2}\[?(?:COPY|TEXTO) OVERLAY\]?\*{0,2}:?\*{0,2}\s*', '', stripped).strip()
            add_texto_overlay_line(doc, text)
            i += 1
            continue

        # SEGUNDO X-Y | Section headers
        if re.match(r'^\*?\*?(?:SEGUNDO|\[?\d+s)', stripped):
            section_text = stripped.replace('**', '').strip()
            add_section_subheader(doc, section_text)
            i += 1
            continue

        # Notas de grabación header (standalone)
        if 'NOTAS DE GRABACION' in stripped.upper() or 'NOTAS DE GRABACIÓN' in stripped.upper():
            add_notes_header(doc, stripped.replace('#', '').replace('*', '').strip())
            i += 1
            continue

        # POR QUE FUNCIONA (standalone) — SKIP
        if is_porque_section(stripped):
            in_porque = True
            i += 1
            continue

        # GUION # headers inside retargeting/level files
        if re.match(r'^GUION\s*#?\d', stripped) or stripped.startswith('**GUION'):
            guion_text = stripped.replace('**', '').strip()
            anuncio_counter += 1
            add_anuncio_banner(doc, anuncio_counter)
            add_section_subheader(doc, guion_text, bg_hex=GRIS_CLARO_HEX, color=AZUL)
            i += 1
            continue

        # Ficha técnica fields — EXCLUDE client-hidden fields (Plataforma, etc.)
        if ':' in stripped and any(stripped.startswith(prefix) for prefix in FICHA_FIELDS):
            # Skip fields that are internal-only (not for client)
            if any(stripped.startswith(excl) for excl in CLIENT_EXCLUDE_FIELDS):
                i += 1
                continue
            parts = stripped.split(':', 1)
            add_field_line(doc, parts[0].strip(), parts[1].strip() if len(parts) > 1 else "")
            i += 1
            continue

        # Ficha técnica table format
        if stripped.startswith('| **') or (stripped.startswith('|') and 'Campo' in stripped):
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if text.startswith('**') and '**' in text[2:]:
                bold_end = text.index('**', 2)
                bold_text = text[2:bold_end]
                rest_text = text[bold_end+2:]
                run_b = p.add_run("• " + bold_text)
                run_b.bold = True
                run_b.font.size = Pt(10)
                run_b.font.color.rgb = GRIS_OSCURO
                if rest_text:
                    run_r = p.add_run(rest_text)
                    run_r.font.size = Pt(10)
                    run_r.font.color.rgb = GRIS_MEDIO
            else:
                run = p.add_run("• " + text)
                run.font.size = Pt(10)
                run.font.color.rgb = GRIS_OSCURO
            i += 1
            continue

        # Numbered items
        if re.match(r'^\d+\.', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = stripped
            if '**' in text:
                parts = re.split(r'\*\*', text)
                for idx_p, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(10)
                        run.font.color.rgb = GRIS_OSCURO
                        if idx_p % 2 == 1:
                            run.bold = True
            else:
                run = p.add_run(text)
                run.font.size = Pt(10)
                run.font.color.rgb = GRIS_OSCURO
            i += 1
            continue

        # Default: normal text
        text = stripped.replace('**', '')
        if text:
            add_normal_text(doc, text)
        i += 1

    doc.save(docx_path)
    print(f"OK: {docx_path}")


def process_md_to_images_docx(md_path, docx_path):
    """Extract ONLY image creatives from .md into a separate styled .docx.

    Returns True if images were found and file was created, False otherwise.
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Check if file has any image creatives
    if 'IMAGEN CREATIVO' not in content:
        return False

    doc = Document()
    setup_doc_styles(doc)

    lines = content.split('\n')

    # Get the original title
    original_title = ""
    for line in lines:
        s = line.strip()
        if s.startswith('# ') and not s.startswith('## '):
            original_title = s[2:].strip()
            break

    # Document header
    add_styled_heading(doc, "CREATIVOS ESTÁTICOS", level=1, color=AZUL)
    if original_title:
        add_normal_text(doc, original_title)
    add_thick_line(doc, TURQUESA_HEX)

    in_image_creative = False
    creativo_count = 0
    i = 0

    while i < len(lines):
        stripped = lines[i].strip()

        # Skip empties, separators, code fences
        if not stripped or stripped in ['---', '```'] or stripped.startswith('===='):
            i += 1
            continue

        # Image section headers (### Imágenes Nivel ...)
        if stripped.startswith('###') and ('Imágenes' in stripped or 'Imagenes' in stripped):
            subheader = stripped.lstrip('#').strip()
            add_thick_line(doc, AZUL_HEX)
            add_section_subheader(doc, subheader, bg_hex=TURQUESA_CLARO_HEX, color=NEGRO)
            in_image_creative = False
            i += 1
            continue

        # CREATIVOS ESTATICOS main headers (## or #)
        if stripped.startswith('#') and ('CREATIVOS' in stripped.upper() or 'ESTATICOS' in stripped.upper() or 'ESTÁTICOS' in stripped):
            text = stripped.lstrip('#').strip()
            add_thick_line(doc, AZUL_HEX)
            add_level_banner(doc, text, AZUL_HEX, BLANCO)
            in_image_creative = False
            i += 1
            continue

        # IMAGEN CREATIVO block start
        if 'IMAGEN CREATIVO' in stripped:
            clean_text = stripped.replace('**', '').strip()
            creativo_count += 1
            add_image_creative_header(doc, clean_text)
            in_image_creative = True
            i += 1
            continue

        # Fields inside an image creative block
        if in_image_creative and ':' in stripped:
            is_img_field = any(stripped.replace('"', '').startswith(f) for f in IMAGE_FIELDS)
            if is_img_field:
                parts = stripped.split(':', 1)
                add_field_line(doc, parts[0].strip(), parts[1].strip() if len(parts) > 1 else "")
                i += 1
                continue
            else:
                # Not an image field — might be leaving the block
                in_image_creative = False

        # Skip everything that's not image-related
        i += 1

    if creativo_count > 0:
        doc.save(docx_path)
        print(f"OK (images): {docx_path}")
        return True
    return False


def get_images_docx_name(md_filename):
    """Generate the IMAGENES_CREATIVOS filename from the original .md filename."""
    base = md_filename.replace('.md', '')
    if base.startswith('GUIONES_'):
        return base.replace('GUIONES_', 'IMAGENES_CREATIVOS_', 1) + '.docx'
    else:
        return 'IMAGENES_CREATIVOS_' + base + '.docx'


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 generate_docx.py <input.md> <output.docx>")
        print("       python3 generate_docx.py --all <base_dir>")
        sys.exit(1)

    if sys.argv[1] == '--all':
        base_dir = sys.argv[2]

        # Process subfolders
        for folder_name in sorted(os.listdir(base_dir)):
            folder_path = os.path.join(base_dir, folder_name)
            if os.path.isdir(folder_path):
                for fname in os.listdir(folder_path):
                    if fname.endswith('.md'):
                        md_path = os.path.join(folder_path, fname)
                        docx_path = os.path.join(folder_path, fname.replace('.md', '.docx'))
                        try:
                            process_md_to_docx(md_path, docx_path)
                        except Exception as e:
                            print(f"ERROR: {md_path}: {e}")
                        # Also generate separate images .docx
                        images_docx = os.path.join(folder_path, get_images_docx_name(fname))
                        try:
                            process_md_to_images_docx(md_path, images_docx)
                        except Exception as e:
                            print(f"ERROR (images): {md_path}: {e}")

        # Process root-level .md files
        for fname in os.listdir(base_dir):
            if fname.endswith('.md') and os.path.isfile(os.path.join(base_dir, fname)):
                md_path = os.path.join(base_dir, fname)
                docx_path = os.path.join(base_dir, fname.replace('.md', '.docx'))
                try:
                    process_md_to_docx(md_path, docx_path)
                except Exception as e:
                    print(f"ERROR: {md_path}: {e}")
                images_docx = os.path.join(base_dir, get_images_docx_name(fname))
                try:
                    process_md_to_images_docx(md_path, images_docx)
                except Exception as e:
                    print(f"ERROR (images): {md_path}: {e}")
    else:
        md_path = sys.argv[1]
        docx_path = sys.argv[2]
        process_md_to_docx(md_path, docx_path)
        # Also generate images docx alongside
        images_docx = os.path.join(
            os.path.dirname(docx_path),
            get_images_docx_name(os.path.basename(md_path))
        )
        process_md_to_images_docx(md_path, images_docx)


if __name__ == '__main__':
    main()
